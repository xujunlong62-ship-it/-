# -*- coding: utf-8 -*-
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

load_dotenv()

client = OpenAI(
    api_key=os.getenv("dashscope_api_key"),
    base_url=os.getenv("dashscope_base_url", "https://dashscope.aliyuncs.com/compatible-mode/v1"),
)

def ai(prompt):
    return client.chat.completions.create(
        model="qwen-plus",
        messages=[{"role": "user", "content": prompt}]
    ).choices[0].message.content

def collect_info():
    print("=== 简历生成器 ===")
    print("请依次输入以下信息（输入空行表示跳过该字段）：")
    print()
    
    info = {}
    
    fields = [
        ("name", "姓名"),
        ("phone", "电话号码"),
        ("email", "邮箱"),
        ("position", "目标职位"),
        ("school", "学校"),
        ("major", "专业"),
        ("degree", "学历（如本科/硕士）"),
        ("grad_year", "毕业时间"),
        ("skills", "技能（用逗号分隔，如 Python, Java, SQL）"),
    ]
    
    for key, label in fields:
        val = input(f"{label}: ").strip()
        info[key] = val
    
    print()
    print("请告诉我你的工作经历（每段经历单独描述，输入完成后输入\"结束\"）：")
    work_exps = []
    while True:
        exp = input("  工作经历 > ").strip()
        if exp == "结束" or exp == "":
            break
        work_exps.append(exp)
    info["work_exps"] = work_exps
    
    print()
    print("请告诉我你的项目经历（每个项目单独描述，输入完成后输入\"结束\"）：")
    projects = []
    while True:
        proj = input("  项目经历 > ").strip()
        if proj == "结束" or proj == "":
            break
        projects.append(proj)
    info["projects"] = projects
    
    return info

def build_resume_text(info):
    parts = []
    parts.append("[PHOTO:150x200]")
    parts.append("[SECTION:个人信息] [TEXT:{} | {} | {} | {}]".format(
        info["name"], info["position"], info["phone"], info["email"]
    ))
    
    if info["school"]:
        parts.append("[SECTION:教育背景] [TEXT:{} | {} | {} | {}]".format(
            info["school"], info["major"], info["degree"], info["grad_year"]
        ))
    
    if info["skills"]:
        for skill in info["skills"].split(","):
            skill = skill.strip()
            if skill:
                parts.append("[SECTION:技能] [ITEM:{}]".format(skill))
    
    if info["work_exps"]:
        for exp in info["work_exps"]:
            parts.append("[SECTION:工作经历] [ITEM:{}]".format(exp))
    
    if info["projects"]:
        for proj in info["projects"]:
            parts.append("[SECTION:项目经历] [ITEM:{}]".format(proj))
    
    return "\n".join(parts)

def get_demo_info():
    return {
        "name": "张三",
        "phone": "138-0000-0000",
        "email": "zhangsan@email.com",
        "position": "Java 后端开发工程师",
        "school": "复旦大学",
        "major": "计算机科学与技术",
        "degree": "本科",
        "grad_year": "2024年6月",
        "skills": "Python, Java, SQL, Docker, Spring Boot, Redis, MySQL, Git, Linux, Vue3",
        "work_exps": [
            "在ABC科技有限公司担任后端开发工程师，使用Spring Boot + MyBatis开发用户管理系统，日活跃用户5万+",
            "优化核心API接口性能，将平均响应时间从800ms降低至150ms，提升81%",
            "设计并实施基于Redis的多级缓存方案，减少60%数据库查询压力，年省服务器成本20万+",
            "主导微服务拆分重构，将单体应用拆分为8个独立服务，支持独立部署与扩容"
        ],
        "projects": [
            "智能排课系统 | Vue3 + Spring Cloud微服务架构 | 遗传算法实现排课优化，效率提升300%，已服务2000+师生",
            "电商订单系统 | 高并发场景下订单处理 | 采用RocketMQ异步解耦，峰值QPS 5000+，数据最终一致性保障",
            "分布式任务调度平台 | XXL-JOB二次开发 | 支持动态分片、失败重试、告警通知，日均调度任务10万+"
        ]
    }

def generate_resume(info):
    text = build_resume_text(info)
    p = ("你是简历顾问。根据以下信息生成专业简历。\n"
         "输出格式：用 [SECTION:标题名] 表示区块标题，用 [ITEM:] 表示列表项，用 [TEXT:] 表示普通文本，用 [PHOTO:尺寸] 表示照片占位。\n"
         "第一行必须是：[PHOTO:150x200]\n"
         "然后是：[SECTION:个人信息] [TEXT:姓名 | 职位 | 联系方式]\n"
         "[SECTION:教育背景] [TEXT:学校 | 专业 | 学历 | 时间]\n"
         "[SECTION:技能] [ITEM:技能1] [ITEM:技能2]\n"
         "[SECTION:工作经历] [SECTION:公司 | 职位 | 时间] [ITEM:经历1]\n"
         "[SECTION:项目经历] [SECTION:项目名] [ITEM:描述]\n"
         "要求：STAR法则、量化成果、精炼、没有的信息省略、不要多余解释。\n"
         "个人信息：" + text)
    return ai(p)

def polish_resume(text):
    p = ("润色这份简历：优化表达、强化动词、突出成果、删除冗余。保持相同格式输出。\n"
         "原始简历：\n" + text)
    return ai(p)

def generate_interview_questions(resume_text, jd_text=None):
    p = ("你是面试官。根据以下简历内容，生成8个高频面试问题。\n"
         "要求：问题由浅入深，包含技术问题和行为问题。\n"
         "输出格式：每个问题前加 [Q:] 前缀。\n"
         "简历内容：\n" + resume_text)
    return ai(p)

# Color constants
COLOR_PRIMARY = "2E5090"
COLOR_ACCENT = "3A6EA5"
COLOR_HEADER = "1A1A1A"
COLOR_BODY = "444444"
COLOR_DIVIDER = "D0D0D0"
COLOR_WHITE = "FFFFFF"
FONT_PRIMARY = "微软雅黑"


def optimize_resume(resume_text, jd_text):
    prompt = f"""你是资深简历优化专家。根据职位描述(JD)优化简历，使简历更匹配该职位。

职位描述：
{jd_text}

原始简历：
{resume_text}

请输出优化后的完整简历，保持原有格式。
"""
    return ai(prompt)


def score_match(resume_text, jd_text):
    prompt = f"""你是招聘系统匹配算法。分析以下简历与职位描述(JD)的匹配程度，只输出0-100的数字。

简历：
{resume_text}

职位描述：
{jd_text}

匹配分数（只输出数字）："""
    try:
        result = ai(prompt).strip()
        return int(result)
    except:
        return 60


FONT_SECONDARY = "Calibri"

def set_font(run, size=10, bold=False, color="333333", name=FONT_PRIMARY):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    for child in list(rPr):
        if child.tag.endswith("rFonts"):
            rPr.remove(child)
    rPr.insert(0, rFonts)

def add_section_header(cell, title):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    set_font(run, size=10, bold=True, color="1A1A2E", name=FONT_PRIMARY)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E5090")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body_text(cell, text, size=9.5, bold=False, space_after=4, indent=0):
    p = cell.add_paragraph(text)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for run in p.runs:
        set_font(run, size=size, bold=bold, color="4A5568", name=FONT_PRIMARY)

def add_bullet(cell, text, size=9.5, space_after=4, indent=0.4):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(indent)
    run_b = p.add_run("• ")
    set_font(run_b, size=size, color="2E5090", name=FONT_PRIMARY)
    run_t = p.add_run(text)
    set_font(run_t, size=size, color="4A5568", name=FONT_PRIMARY)

def save_word(text, filename):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    tokens = re.findall(r"\[(?:SECTION|ITEM|TEXT|PHOTO):[^\]]*\]", text)
    sections = {}
    current_section = None
    for token in tokens:
        inner = token[1:-1]
        if ":" in inner:
            token_type, token_value = inner.split(":", 1)
        else:
            continue
        if token_type == "SECTION":
            current_section = token_value
            if current_section not in sections:
                sections[current_section] = []
        elif token_type == "ITEM":
            if current_section:
                sections[current_section].append(("item", token_value))
        elif token_type == "TEXT":
            if current_section:
                sections[current_section].append(("text", token_value))
    if not sections:
        sections = {"简历": [("text", text.strip())]}

    personal_info = ""
    if "个人信息" in sections:
        for typ, val in sections["个人信息"]:
            if typ == "text":
                personal_info = val

    name = ""
    position = ""
    phone = ""
    email = ""
    if personal_info:
        parts = [p.strip() for p in personal_info.split("|")]
        if len(parts) >= 1: name = parts[0]
        if len(parts) >= 2: position = parts[1]
        if len(parts) >= 3: phone = parts[2]
        if len(parts) >= 4: email = parts[3]

    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = True
    left_h = header_tbl.cell(0, 0)
    right_h = header_tbl.cell(0, 1)
    left_h.width = Cm(12)
    right_h.width = Cm(7.2)

    for cell in [left_h, right_h]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement("w:" + bn)
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    p_name = left_h.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(1)
    run_name = p_name.add_run(name or "姓名")
    set_font(run_name, size=26, bold=True, color="1A1A2E", name=FONT_PRIMARY)

    if position:
        p_pos = left_h.add_paragraph()
        p_pos.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pos.paragraph_format.space_before = Pt(0)
        p_pos.paragraph_format.space_after = Pt(6)
        run_pos = p_pos.add_run(position)
        set_font(run_pos, size=12, color="4A5568", name=FONT_PRIMARY)

    contact_parts = []
    if phone: contact_parts.append(phone)
    if email: contact_parts.append(email)
    if contact_parts:
        p_contact = left_h.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(0)
        run_contact = p_contact.add_run("  •  ".join(contact_parts))
        set_font(run_contact, size=9.5, color="718096", name=FONT_PRIMARY)

    p_photo = right_h.add_paragraph()
    p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_photo.paragraph_format.space_before = Pt(6)
    p_photo.paragraph_format.space_after = Pt(2)
    run_photo = p_photo.add_run("[ 照片 ]")
    set_font(run_photo, size=10, color="A0AEC0", name=FONT_PRIMARY)

    p_ph_box = right_h.add_paragraph()
    p_ph_box.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ph_box.paragraph_format.space_before = Pt(0)
    p_ph_box.paragraph_format.space_after = Pt(0)
    run_box = p_ph_box.add_run("━" * 16 + "\n" + "━" * 16)
    set_font(run_box, size=8, color="E2E8F0")

    header_tbl.rows[0].height = Cm(2.8)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(8)
    run_line = p_line.add_run("━" * 72)
    set_font(run_line, size=4, color="E2E8F0")

    col_left_w = Cm(7.0)
    col_right_w = Cm(11.5)

    main_tbl = doc.add_table(rows=1, cols=2)
    main_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_tbl.autofit = True
    left = main_tbl.cell(0, 0)
    right = main_tbl.cell(0, 1)
    left.width = col_left_w
    right.width = col_right_w

    for cell in [left, right]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for bn in ["top", "left", "bottom", "right"]:
            b = OxmlElement("w:" + bn)
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    if "教育背景" in sections:
        add_section_header(left, "教育背景")
        for typ, val in sections["教育背景"]:
            add_body_text(left, val, size=9.5, space_after=4)

    if "技能" in sections:
        add_section_header(left, "技能")
        for typ, val in sections["技能"]:
            add_bullet(left, val, size=9, space_after=2, indent=0.3)

    if "工作经历" in sections:
        add_section_header(right, "工作经历")
        for typ, val in sections["工作经历"]:
            add_bullet(right, val, size=9.5, space_after=6, indent=0.4)

    if "项目经历" in sections:
        add_section_header(right, "项目经历")
        for typ, val in sections["项目经历"]:
            add_bullet(right, val, size=9.5, space_after=6, indent=0.4)

    base_name = filename.rsplit(".docx", 1)[0]
    counter = 1
    while os.path.exists(filename):
        filename = base_name + "_" + str(counter) + ".docx"
        counter += 1
    doc.save(filename)
    print("简历已保存: " + filename)

def save_questions(text, filename):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p = doc.add_paragraph()
    run = p.add_run("面试题库")
    set_font(run, size=20, bold=True, color="2E5090", name=FONT_PRIMARY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("━" * 40)
    set_font(run, size=8, color="2E5090")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    q_num = 0
    for line in lines:
        if line.startswith("[Q:"):
            q_num += 1
            q_text = line[3:].rstrip("]") if line.endswith("]") else line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run_q = p.add_run(str(q_num) + ". ")
            set_font(run_q, size=11, bold=True, color="2E5090", name=FONT_PRIMARY)
            run_a = p.add_run(q_text)
            set_font(run_a, size=10.5, color="4A5568", name=FONT_PRIMARY)
        elif line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_font(run, size=10, color="4A5568", name=FONT_PRIMARY)

    doc.save(filename)
    print("面试题库已保存: " + filename)

def main():
    import sys
    use_demo = False
    
    if len(sys.argv) > 1 and sys.argv[1] in ("--demo", "-d", "demo"):
        use_demo = True
    else:
        print()
        print("=" * 40)
        print("       欢迎使用简历生成器")
        print("=" * 40)
        print()
        print("请选择模式：")
        print("  1. 交互式输入（手动填写信息）")
        print("  2. 演示模式（使用内置样本数据快速生成）")
        print()
        choice = input("请输入选择 [1/2，默认1]: ").strip()
        if choice == "2":
            use_demo = True
    
    if use_demo:
        print()
        print(">>> 使用演示模式，基于样本数据生成简历...")
        info = get_demo_info()
        resume_text = build_resume_text(info)
    else:
        print()
        print("=" * 40)
        print("       欢迎使用简历生成器")
        print("=" * 40)
        print()
        print("我会依次问你一些基本信息，")
        print("你只需如实回答即可。")
        print("遇到不重要的字段可以直接按回车跳过。")
        print()
        
        info = collect_info()
        print()
        print("正在生成简历...")
        resume_raw = generate_resume(info)
        
        print("正在优化简历...")
        resume_text = polish_resume(resume_raw)
    
    print()
    print("正在生成 Word 文档...")
    save_word(resume_text, "简历.docx")
    
    print()
    print("正在生成面试题库...")
    questions = generate_interview_questions(resume_text)
    save_questions(questions, "简历_面试题库.docx")
    
    print()
    print("=" * 40)
    print("搞定！你现在有一份优化好的简历和面试题库了！")
    print("=" * 40)

if __name__ == "__main__":
    main()