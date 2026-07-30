# -*- coding: utf-8 -*-
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

load_dotenv()

client = OpenAI(
    api_key=os.getenv('dashscope_api_key'),
    base_url=os.getenv('dashscope_base_url', 'https://dashscope.aliyuncs.com/compatible-mode/v1'),
)

def ai(prompt):
    return client.chat.completions.create(
        model='qwen-plus',
        messages=[{'role': 'user', 'content': prompt}]
    ).choices[0].message.content

def collect_info():
    print('=== 简历生成器 ===')
    print('请依次输入以下信息（输入空行表示跳过该字段）：')
    print()
    
    info = {}
    
    # 收集基本信息
    fields = [
        ('name',       '姓名'),
        ('phone',      '电话号码'),
        ('email',      '邮箱'),
        ('position',   '目标职位'),
        ('school',     '学校'),
        ('major',      '专业'),
        ('degree',     '学历（如本科/硕士）'),
        ('grad_year',  '毕业时间'),
        ('skills',     '技能（用逗号分隔，如 Python, Java, SQL）'),
    ]
    
    for key, label in fields:
        val = input(f'{label}: ').strip()
        info[key] = val
    
    # 收集工作经历
    print()
    print('请告诉我你的工作经历（每段经历单独描述，输入完成后输入"结束"）：')
    work_exps = []
    while True:
        exp = input('  工作经历 > ').strip()
        if exp == '结束' or exp == '':
            break
        work_exps.append(exp)
    info['work_exps'] = work_exps
    
    # 收集项目经历
    print()
    print('请告诉我你的项目经历（每个项目单独描述，输入完成后输入"结束"）：')
    projects = []
    while True:
        proj = input('  项目经历 > ').strip()
        if proj == '结束' or proj == '':
            break
        projects.append(proj)
    info['projects'] = projects
    
    return info

def build_resume_text(info):
    parts = []
    parts.append('[PHOTO:150x200]')
    parts.append('[SECTION:个人信息] [TEXT:{} | {} | {} | {}]'.format(
        info['name'], info['position'], info['phone'], info['email']
    ))
    
    if info['school']:
        parts.append('[SECTION:教育背景] [TEXT:{} | {} | {} | {}]'.format(
            info['school'], info['major'], info['degree'], info['grad_year']
        ))
    
    if info['skills']:
        for skill in info['skills'].split('，'):
            skill = skill.strip()
            if skill:
                parts.append('[SECTION:技能] [ITEM:{}]'.format(skill))
    
    if info['work_exps']:
        for exp in info['work_exps']:
            parts.append('[SECTION:工作经历] [ITEM:{}]'.format(exp))
    
    if info['projects']:
        for proj in info['projects']:
            parts.append('[SECTION:项目经历] [ITEM:{}]'.format(proj))
    
    return '\n'.join(parts)

def generate_resume(info):
    text = build_resume_text(info)
    p = ('你是简历顾问。根据以下信息生成专业简历。\n'
         '输出格式：用 [SECTION:标题名] 表示区块标题，用 [ITEM:] 表示列表项，用 [TEXT:] 表示普通文本，用 [PHOTO:尺寸] 表示照片占位。\n'
         '第一行必须是：[PHOTO:150x200]\n'
         '然后是：[SECTION:个人信息] [TEXT:姓名 | 职位 | 联系方式]\n'
         '[SECTION:教育背景] [TEXT:学校 | 专业 | 学历 | 时间]\n'
         '[SECTION:技能] [ITEM:技能1] [ITEM:技能2]\n'
         '[SECTION:工作经历] [SECTION:公司 | 职位 | 时间] [ITEM:经历1]\n'
         '[SECTION:项目经历] [SECTION:项目名] [ITEM:描述]\n'
         '要求：STAR法则、量化成果、精炼、没有的信息省略、不要多余解释。\n'
         '个人信息：' + text)
    return ai(p)

def polish_resume(text):
    p = ('润色这份简历：优化表达、强化动词、突出成果、删除冗余。保持相同格式输出。\n'
         '原始简历：\n' + text)
    return ai(p)

def generate_interview_questions(resume_text):
    p = ('你是面试官。根据以下简历内容，生成8个高频面试问题。\n'
         '要求：问题由浅入深，包含技术问题和行为问题。\n'
         '输出格式：每个问题前加 [Q:] 前缀。\n'
         '简历内容：\n' + resume_text)
    return ai(p)

def set_font(run, size=10, bold=False, color='333333', name='微软雅黑'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    # Remove existing rFonts to avoid conflicts
    for child in list(rPr):
        if child.tag.endswith('rFonts'):
            rPr.remove(child)
    rPr.insert(0, rFonts)

def save_word(text, filename):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # Parse tokens like [SECTION:name], [ITEM:val], [TEXT:val], [PHOTO:...]
    # Handle multiple tokens on the same line
    tokens = re.findall(r'\[(?:SECTION|ITEM|TEXT|PHOTO):[^\]]*\]', text)

    sections = {}
    current_section = None

    for token in tokens:
        inner = token[1:-1]  # Remove surrounding [ ]
        if ":" in inner:
            token_type, token_value = inner.split(":", 1)
        else:
            continue

        if token_type == "SECTION":
            current_section = token_value
            if current_section not in sections:
                sections[current_section] = []
        elif token_type == "ITEM":
            if current_section:
                sections[current_section].append(("item", token_value))
        elif token_type == "TEXT":
            if current_section:
                sections[current_section].append(("text", token_value))
        # PHOTO is ignored
    if not sections:
        sections = {'简历': [('text', text.strip())]}

    personal_info = ''
    if '个人信息' in sections:
        for typ, val in sections['个人信息']:
            if typ == 'text':
                personal_info = val
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(personal_info)
    set_font(run, size=16, bold=True, color='1A1A1A')
    p.paragraph_format.space_after = Pt(2)
    
    p = doc.add_paragraph()
    run = p.add_run('\u2501' * 45)
    set_font(run, size=10, color='2E5090')
    p.paragraph_format.space_after = Pt(6)

    cols_width_left = Cm(7.5)
    cols_width_right = Cm(11.5)
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = cols_width_left
    right_cell.width = cols_width_right
    
    for cell in [left_cell, right_cell]:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for bn in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement('w:' + bn)
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            tcBorders.append(b)
        tcPr.append(tcBorders)
    
    left_content = []
    if '教育背景' in sections:
        left_content.append(('header', '教育背景'))
        for typ, val in sections['教育背景']:
            left_content.append((typ, val))
    if '技能' in sections:
        left_content.append(('header', '技能'))
        for typ, val in sections['技能']:
            left_content.append((typ, val))
    
    for typ, val in left_content:
        if typ == 'header':
            p = left_cell.add_paragraph()
            run = p.add_run(val)
            set_font(run, size=10.5, bold=True, color='2E5090')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p2 = left_cell.add_paragraph()
            run2 = p2.add_run('\u2501' * 12)
            set_font(run2, size=6, color='D0D0D0')
            p2.paragraph_format.space_after = Pt(3)
        elif typ == 'text':
            p = left_cell.add_paragraph(val)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_font(run, size=9.5)
        elif typ == 'item':
            p = left_cell.add_paragraph(val)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            for run in p.runs:
                set_font(run, size=9.5)

    right_content = []
    if '工作经历' in sections:
        right_content.append(('header', '工作经历'))
        for typ, val in sections['工作经历']:
            right_content.append((typ, val))
    if '项目经历' in sections:
        right_content.append(('header', '项目经历'))
        for typ, val in sections['项目经历']:
            right_content.append((typ, val))
    
    for typ, val in right_content:
        if typ == 'header':
            p = right_cell.add_paragraph()
            run = p.add_run(val)
            set_font(run, size=10.5, bold=True, color='2E5090')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p2 = right_cell.add_paragraph()
            run2 = p2.add_run('\u2501' * 18)
            set_font(run2, size=6, color='D0D0D0')
            p2.paragraph_format.space_after = Pt(3)
        elif typ == 'text':
            p = right_cell.add_paragraph(val)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                set_font(run, size=10)
        elif typ == 'item':
            p = right_cell.add_paragraph(val)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.3)
            for run in p.runs:
                set_font(run, size=10)
    
    base_name = filename.rsplit('.docx', 1)[0]
    counter = 1
    while os.path.exists(filename):
        filename = base_name + '_' + str(counter) + '.docx'
        counter += 1
    doc.save(filename)
    print('简历已保存: ' + filename)

def save_questions(text, filename):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    run = p.add_run('面试题库')
    set_font(run, size=18, bold=True, color='1A1A1A')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph()
    run = p.add_run('\u2501' * 45)
    set_font(run, size=10, color='2E5090')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines:
        if line.startswith('[Q:'):
            question_text = line[3:] if line.endswith(']') else line[3:]
            p = doc.add_paragraph()
            run = p.add_run(question_text)
            set_font(run, size=11, bold=True, color='2E5090')
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)

    base_name = filename.rsplit('.docx', 1)[0]
    filename = base_name + '_面试题库' + '.docx'
    counter = 1
    while os.path.exists(filename):
        filename = base_name + '_面试题库_' + str(counter) + '.docx'
        counter += 1
    doc.save(filename)
    print('面试题库已保存: ' + filename)

def main():
    print()
    print('=' * 40)
    print('       欢迎使用简历生成器')
    print('=' * 40)
    print()
    print('我会依次问你一些基本信息，')
    print('你只需如实回答即可。')
    print('遇到不重要的字段可以直接按回车跳过。')
    print()
    
    info = collect_info()
    
    print()
    print('正在生成简历...')
    resume_raw = generate_resume(info)
    
    print('正在优化简历...')
    resume_polished = polish_resume(resume_raw)
    
    print('正在生成 Word 文档...')
    save_word(resume_polished, '简历.docx')
    
    print()
    print('正在生成面试题库...')
    questions = generate_interview_questions(resume_polished)
    save_questions(questions, '简历_面试题库.docx')
    
    print()
    print('=' * 40)
    print('搞定！你现在有一份优化好的简历和面试题库了！')
    print('=' * 40)

if __name__ == '__main__':
    main()
